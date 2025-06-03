# test_word_parser.py
import unittest
import os
from docx import Document
from parsers.word_parser import WordParser


class TestWordParser(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        """创建测试用的Word文档"""
        cls.test_doc_path = "test_references.docx"
        doc = Document()

        # 添加非参考文献内容
        doc.add_heading("正文部分", level=1)
        doc.add_paragraph("这是正文内容")

        # 添加参考文献部分
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 作者1. 文章1. 期刊1, 2020.")
        doc.add_paragraph("[2] 作者2. 文章2. 期刊2, 2021.")
        doc.add_paragraph("")  # 空行
        doc.add_paragraph("其他内容")  # 参考文献后的内容

        doc.save(cls.test_doc_path)

        # 创建无参考文献的文档
        cls.empty_doc_path = "test_empty.docx"
        doc = Document()
        doc.add_paragraph("没有参考文献的文档")
        doc.save(cls.empty_doc_path)

    @classmethod
    def tearDownClass(cls):
        """删除测试文件"""
        os.remove(cls.test_doc_path)
        os.remove(cls.empty_doc_path)

    def test_extract_references_normal(self):
        """测试正常提取参考文献"""
        expected = [
            "[1] 作者1. 文章1. 期刊1, 2020.",
            "[2] 作者2. 文章2. 期刊2, 2021."
        ]
        result = WordParser.extract_references(self.test_doc_path)
        self.assertEqual(result, expected)

    def test_extract_references_no_references(self):
        """测试文档中没有参考文献的情况"""
        result = WordParser.extract_references(self.empty_doc_path)
        self.assertEqual(result, [])

    def test_extract_references_empty_doc(self):
        """测试空文档的情况"""
        empty_doc_path = "test_empty_doc.docx"
        doc = Document()
        doc.save(empty_doc_path)

        result = WordParser.extract_references(empty_doc_path)
        self.assertEqual(result, [])

        os.remove(empty_doc_path)

    def test_extract_references_multiple_sections(self):
        """测试文档中有多个'参考文献'标题的情况"""
        multi_doc_path = "test_multi_references.docx"
        doc = Document()

        # 第一个参考文献部分
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 文献1")

        # 其他内容
        doc.add_paragraph("中间内容")

        # 第二个参考文献部分
        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[2] 文献2")

        doc.save(multi_doc_path)

        # 应该只提取第一个参考文献部分
        expected = ["[1] 文献1"]
        result = WordParser.extract_references(multi_doc_path)
        self.assertEqual(result, expected)

        os.remove(multi_doc_path)

    def test_extract_references_no_empty_line(self):
        """测试参考文献部分没有空行结束的情况"""
        no_empty_doc_path = "test_no_empty_line.docx"
        doc = Document()

        doc.add_heading("参考文献", level=1)
        doc.add_paragraph("[1] 文献1")
        doc.add_paragraph("[2] 文献2")
        # 没有空行直接结束

        doc.save(no_empty_doc_path)

        expected = ["[1] 文献1", "[2] 文献2"]
        result = WordParser.extract_references(no_empty_doc_path)
        self.assertEqual(result, expected)

        os.remove(no_empty_doc_path)


if __name__ == '__main__':
    unittest.main()
